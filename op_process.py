from deepgram import DeepgramClient, PrerecordedOptions
import json
from chat import *
from docx import Document

def process_json_op(op):
    opdata = json.loads(op)
    description = opdata['results']['channels']
    print(description[0]['alternatives'][0]['paragraphs']['paragraphs'])
    text_transcript = {}
   
    for i in range(0,len(description[0]['alternatives'][0]['paragraphs']['paragraphs'])):
    # for i in range(0,1):
        
        text_combined =''
        inter = description[0]['alternatives'][0]['paragraphs']['paragraphs'][i]['sentences']
        start = inter[0]['start']
        end = inter[0]['end']
        # print("Inter", inter)
        for item in inter:
            # Append the text value to the combined text string
            text_combined += item['text'] + " "
            # Update start and end with the minimum and maximum times
            start = min(start, item['start'])
            end = max(end, item['end'])
        
        text_transcript[f'sentence_{i}'] = {
            'text': text_combined.strip(), # Remove trailing space
            'start': start,
            'end': end
        }

    return text_transcript


def doc(summary,op_file):
# Your string
    string = f'{summary}'

    # Create a new Word document
    doc = Document()

    # Split the string into sections
    sections = string.split('\n\n')

    for section in sections:
        # Split each section into lines
        lines = section.split('\n')
        
        # Add the section title as a heading
        doc.add_heading(lines[0].replace('**', ''), level=1)
        
        # Add the timestamp and details as paragraphs
        for line in lines[1:]:
            if line.startswith('- **Timestamp**'):
                doc.add_paragraph(line, style='ListBullet')
            elif line.startswith('- **Details**'):
                doc.add_paragraph(line, style='ListBullet')
            else:
                doc.add_paragraph(line)

    # Save the document
    doc.save(f'{op_file}.docx')