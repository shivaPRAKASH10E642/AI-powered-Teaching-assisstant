from youtube_transcript_api import YouTubeTranscriptApi
import requests
import pandas as pd
from docx import Document

def transcript_generator(video_id):
    transcript_inter = ""
    # Fetch the transcript
    transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['en', 'en-GB'])
    transcript
    # Print the transcript
    for entry in transcript:
        transcript_inter= transcript_inter + entry['text']
        # print(transcript.join(entry['text']))
    # print(f"printing transcript inter {video_id}\n",transcript_inter)
    return transcript_inter


def create_document(text,file,heading):
    # Create a new Document object
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(text)
    document.save(f'{file}.docx')